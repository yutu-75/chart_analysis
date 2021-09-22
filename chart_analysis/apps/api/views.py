from django.shortcuts import render

# Create your views here.
# 引入APIView,APIView是继承的django的View，也就是APIView在View的基础上添加了一些其他的功能
from rest_framework.views import APIView
from rest_framework.response import Response
from django.http import StreamingHttpResponse
from django.utils.encoding import escape_uri_path
from datetime import datetime
from rest_framework import serializers
from rest_framework.parsers import FileUploadParser
import docx
import re
import os
import json
import uuid
import textract
from django.db import connection, connections
from chart_analysis.settings import STATIC_URL_TEMPLATE, MEDIA_ROOT
from . import models
import numpy as np
import uuid
import PyPDF2
import pdfplumber
import decimal

import docx2txt

def content_analysis(str_content,crux_word):
    if not crux_word:
        crux_word = ['奖励', '嘉奖', '补贴', '补助', '补偿', '兜底', '财政支持', '减免', '资助', '税费', '税收', '纳税', '纳税额', '缴纳入库地方级税收',
                     '纳税地方贡献额留成', '纳税地方留成', '地方税收贡献额', '税收地方留成', '地方财政留成', '财政贡献地方留成', '实际形成地方财力', '当地财力贡献', '当地财力实际贡献',
                     '项目库', '名录库', '备选库', '侯选库', '资格库', '储备库', '培育库', '目录库', '供应商库', '中介机构库', '代理机构库', '承包商库', '预选库',
                     '目录', '黑名单', '红名单', '排行榜', '排名榜', '注册', '登记', '备案', '年检', '年报', '监制', '认定', '认可', '检验', '监测', '审定',
                     '指定', '配号', '复检', '复审', '换证', '设立分支机构', '设立法人机构', '设立子公司', '设立分公司', '设有分支机构', '设有子公司', '设有分公司',
                     '常驻机构', '服务网点', '售后机构', '售后服务机构', '迁出', '迁离', '外迁', '搬离', '搬迁', '转出', '追回', '返还', '退还', '收回', '追缴',
                     '承诺', '推荐', '优先给予', '优先支持', '优先扶持', '优先安排', '优先推荐', '优先采购', '优先选用', '优先使用', '优先考虑', '优先选择', '优先选购',
                     '重点支持', '重点扶持', '重点培育', '政策支持', '领军企业', '龙头企业', '知名', '骨干', '国有企业', '民营企业', '限上企业', '规上企业', '本土企业',
                     '本地企业', '本土的企业', '本地的企业', '当地企业', '当地的企业', '外地企业', '外地机构', '区内', '市内', '县内', '本地经营业绩', '本地项目经验',
                     '前置条件', '前置程序', '参考价格', '价格补贴', '统一价格', '本地经营面积', '本地办公面积', '本地固定场所', '本地缴纳社会保险', '减免行政事业性收费',
                     '减免公积金', '减免政府性基金', '缓征行政事业性收费', '缓征公积金', '缓征政府性基金', '停征行政事业性收费', '停征公积金', '停征政府性基金', '减免土地出让金',
                     '减免社会保险费用', '缓征社会保险费用', '缴纳保证金', '以现金形式']

    start_time = datetime.now()
    n = 1
    data_content = []
    serial_list = []
    mapping_list = []
    str_content_text = str_content.split('\n')
    str_content_text = [s_i.strip() for s_i in str_content_text if s_i]

    for i in str_content_text:  # 循环文本的每一行的数据
        i = i.strip()
        if len(i) < 1:
            continue
        n += 1

        def m_i(q, i_data):
            # print(q,i_data)
            if q in mapping_list:
                # print('qwq')

                a_index = mapping_list.index(q)
                mapping_list[:a_index - 1].append(q)
                del serial_list[a_index:]
                serial_list.append(i_data)
            else:
                mapping_list.append(q)
                serial_list.append(i_data)
            # print(serial_list)
            # print()

        i_sign = ""
        if re.findall('^(第.*?章) ', i):  # 匹配    第一章
            # print(i)
            m_i('A', i)
            i_sign = "A"
        elif re.findall('^(第.*?条) ', i):  # 匹配    第一条
            # print(i)
            m_i('B', i)
            i_sign = "B"
        elif re.findall('^([一二三四五六七八九十]+、)', i):  # 匹配     一、
            # print(i)
            m_i('C', i)
            i_sign = "C"
        elif re.findall('^(（[一二三四五六七八九十]+）)', i):  # 匹配    （一）
            # print(i)
            m_i('D', i)
            i_sign = "D"
        elif re.findall('^(\d+\.)', i):  # 匹配     1.
            # print(i)
            m_i('E', i)
            i_sign = "E"
        elif re.findall('^(（\d+）)', i):  # 匹配    （1）
            # print(i)
            m_i('F', i)
            i_sign = "F"
        for c_word in crux_word:
            # if c_word in i:
            sum_i = i.count(c_word)
            if sum_i:

                # print(serial_list,[c_word]+serial_list+[i])
                if i not in serial_list:
                    data_content.append([c_word] + serial_list + [i] + [str(n)] + [str(sum_i)])
                else:
                    data_content.append([c_word] + serial_list + [str(n)] + [str(sum_i)])
                # print(data_content[-1][-3], '--------------', str_content_text.index(i))

                # 段尾获取
                end_text = ''

                for d_c in str_content_text[str_content_text.index(i):]:  # 从当前段落往后循环

                    i_sign1 = ""
                    if re.findall('^(第.*?章) ', d_c):  # 匹配    第一章
                        i_sign1 = "A"
                    elif re.findall('^(第.*?条) ', d_c):  # 匹配    第一条
                        i_sign1 = "B"
                    elif re.findall('^([一二三四五六七八九十]+)', d_c):  # 匹配     一、

                        i_sign1 = "C"
                    elif re.findall('^(（[一二三四五六七八九十]+）)', d_c):  # 匹配    （一）
                        i_sign1 = "D"
                    elif re.findall('^(\d+\.)', d_c):  # 匹配     1.

                        i_sign1 = "E"
                    elif re.findall('^(（\d+）)', d_c):  # 匹配    （1）
                        i_sign1 = "F"
                    else:
                        i_sign1 = "G"
                    # print(mapping_list, i_sign, i_sign1)

                    if i_sign1 == 'G':
                        # end_text = d_c
                        data_content[-1].append(d_c)
                        break
                    if i_sign1 != i_sign:
                        break
    for q in data_content:
        print(q)

    print(datetime.now() - start_time)
    return data_content


def insert_db(data,f_id,f_url,f_name):
    start_time = datetime.now()
    print(start_time)
    with connections['gpjz'].cursor() as cursor:
        for i in data:
            try:
                if i[-1].isalnum():

                    sql = r"""insert into content_analysis_upload(content_id,
                        content_title,
                        content_download_url,
                        word_crux,
                        word_count,
                        word_page,
                        content_appear,
                        content_first,
                        content_end     
                    ) values ("{}","{}","{}","{}",{},"{}","{}","{}","{}")""".format(f_id, f_name, f_url, i[0], i[-1], i[-2], i[-3], i[1:-3], '0')
                    print(sql)
                else:
                    sql = r"""insert into content_analysis_upload(
                        content_id,
                        content_title,
                        content_download_url,
                        word_crux,
                        word_count,
                        word_page,
                        content_appear,
                        content_first,
                        content_end) values ("{}","{}","{}","{}",{},"{}","{}","{}","{}")""".format(f_id, f_name, f_url, i[0], i[-2], i[-3], i[-4], i[1:-4], i[-1])
                    print(sql)

                # 执行sql语句
                cursor.execute(sql)
                # 提交到数据库执行
                connections['gpjz'].commit()
            except Exception as e:
                # 如果发生错误则回滚
                connections['gpjz'].rollback()
                print(e,'这里报错啦！',i)
                break
        print(datetime.now()-start_time)


class Downloads(APIView):
    def get(self, request):
        data = {
            "error_code": 400,
            "params": {
                "error": ""
            }
        }
        id = request.GET.get('id')
        print(id)
        if not id:
            data['params']['error'] = '接口参数缺失！'
            return Response(data=data)
        a = models.File.objects.filter(name=id).first()
        print(a)

        if not a:
            data['error_code'] = 401
            data['params']['error'] = '接口参数错误！'
            return Response(data=data)
        f = a.file
        f_name = str(f).replace('file/','')
        file_path = MEDIA_ROOT.joinpath(str(f))
        file = open(str(file_path), 'rb')
        response =StreamingHttpResponse(file)
        response['Content-Type'] = 'application/octet-stream'
        response['Content-Disposition'] = 'attachment;filename="{}"'.format(escape_uri_path(f_name))  # 这是文件的简单描述，注意写法就是这个固定的写法
        return response


class WordCountView(APIView):
    def get(self, request):

        return Response('no get!!!')

    def post(self, request):
        url_ip = 'http://49.4.31.249/'
        data = {
            "error_code": 400,
            "params": {
                "error": ""
            }
        }
        file_obj = request.FILES.get('file')
        print('file_obj:', file_obj)

        crux_word = request.POST.get('crux_word')
        if crux_word:

            try:
                crux_word = eval(crux_word)
            except:
                crux_word = crux_word.split(',')

        if not file_obj:
            data['params']['error'] = '接口参数缺失！'

            return Response(data=data)

        name_file = models.File.objects.create(
            name=uuid.uuid1(),
            file=file_obj,
        )  # 自动就会将文件上传到我们配置的img文件夹中
        c = name_file.file                  # c: file/南平市人民政府关于印发南平市进一步促进总部经济发展若干措施的通知_M1f3Dz4.docx
        print('c:',c)
        file_path = MEDIA_ROOT.joinpath(str(c))  # 保存文件的路径
        if str(file_obj).split('.')[1] == 'pdf':
            w_content = self.pdf_get(file_path)

        elif str(file_obj).split('.')[1] == 'docx':
            w_content = self.word_get(file_path)

        elif str(file_obj).split('.')[1] == 'doc':

            w_content = self.word_get_doc(file_path)
        else:
            data['error_code'] = 402
            data['params']['error'] = '文件格式错误！拒绝上传！'
            os.remove(file_path)
            models.File.objects.filter(name=name_file.name).delete()
            return Response(data=data, status=402)
        if w_content == 'error':
            data1 = {
                "error_code": 403,
                "params": {
                    "error": "文件可能被加密,损坏，或为空等未知情况无法打开！"
                }
            }
            return Response(data=data1)
        data = content_analysis(w_content, crux_word)


        # insert_db(data,name_file.name,'{}api/downloads/?id={}'.format(url_ip, name_file.name), str(file_obj))
        print("name_file.name:", name_file.name)
        data_list = []

        for d_i in data:
            print(d_i)
            dict_data = {
                'content_id': str(name_file.name),
                'content_title': str(file_obj),
                'content_download_url': '{}api/downloads/?id={}'.format(url_ip, name_file.name),
                'word_crux': d_i[0]
            }

            if d_i[-1].isalnum():

                dict_data['word_count'] = d_i[-1]
                dict_data['word_page'] = d_i[-2]
                dict_data['content_appear'] = d_i[-3]
                dict_data['content_first'] = d_i[1:-3]
                dict_data['content_end'] = '0'
            else:
                dict_data['word_count'] = d_i[-2]
                dict_data['word_page'] = d_i[-3]
                dict_data['content_appear'] = d_i[-4]
                dict_data['content_first'] = d_i[1:-4]
                dict_data['content_end'] = d_i[-1]

            data_list.append(dict_data)

        content_data = {
            "error_code": 200,
            "data": data_list

        }

        return Response(status=200, data=content_data, content_type=json)

    def pdf_get(self, pdf_path):
        ee_time = datetime.now()
        d_data = {}
        page_content = ''
        # 内容提取，使用 pdfplumber 打开 PDF，用于提取文本
        try:
            with pdfplumber.open(pdf_path) as pdf_file:
                # 使用 PyPDF2 打开 PDF 用于提取图片

                pdf_image_reader = PyPDF2.PdfFileReader(open(pdf_path, "rb"))
                d_data['sum_page'] = pdf_image_reader.getNumPages()  # 获取总页数
                print("开始执行---")
                # len(pdf.pages)为PDF文档页数，一页页解析
                for i in range(len(pdf_file.pages)):

                    page_text = pdf_file.pages[i].extract_text()
                    if not page_text:
                        continue
                    # page.extract_text()函数即读取文本内容
                    page_content += page_text
        except Exception as e:
            return 'error'

        print("文本内容获取所用时间:", datetime.now() - ee_time)
        return page_content

    def word_get_doc(self, filename):
        """
        # 此函数方法只能在Linux上运行
        并安装 antiword
        命令:
            apt-get install antiword
        """
        try:
            text = textract.process(filename)
            text = text.decode()
            return text
        except Exception as e:
            print(e)
            return 'error'


    def word_get(self,filename):
        print(filename)
        # document = docx2txt.process(filename)
        try:
            document = docx.Document(filename)
        except Exception as e:
            print(e)
            return 'error'

        print(len(document.paragraphs))                     # 段落长度
        content = []
        for i in range(len(document.paragraphs)):
            content.append(document.paragraphs[i].text)

        tables = document.tables  # 获取文件中的表格集
        qwq = []  # word文档里面表格内容
        for t_i in tables:
            for r_j in t_i.rows:
                for z in r_j.cells:
                    if qwq:
                        if z.text not in qwq:
                            qwq.append(z.text)
                    else:
                        qwq.append(z.text)

        word_content = '\n'.join(qwq + content)
        return word_content