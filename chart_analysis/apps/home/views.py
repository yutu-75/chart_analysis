# coding:utf-8
import json
import os
import jieba
import requests
from django.shortcuts import render
from django.db import connection, connections

from django.http import HttpResponse, JsonResponse
from django.http import FileResponse
from . import models
from datetime import datetime
import datetime as dt
# from pathlib import Path
import random
# Create your views here.
from collections import Counter
import jieba.posseg
from jinja2 import Environment, FileSystemLoader
from pyecharts.charts import Bar, Line, Pie, Scatter, WordCloud, Radar, Funnel, Gauge, Map, Geo
from pyecharts.globals import ChartType
from pyecharts import options as opts
from pyecharts.globals import CurrentConfig
from chart_analysis.settings import STATIC_URL_TEMPLATE, MEDIA_ROOT
from chart_analysis.settings import database_name  # 数据库中文名
# 百分比需要的参数
from pyecharts.commons.utils import JsCode
from pyecharts.globals import ThemeType
import numpy as np
import uuid
import PyPDF2
import pdfplumber
import decimal
from docx import Document

CurrentConfig.GLOBAL_ENV = Environment(loader=FileSystemLoader(r"chart_analysis/apps/home/templates"))

# 引入APIView,APIView是继承的django的View，也就是APIView在View的基础上添加了一些其他的功能
from rest_framework.views import APIView
from rest_framework.response import Response
from django.http import StreamingHttpResponse
from django.utils.encoding import escape_uri_path
from rest_framework import serializers
from rest_framework.parsers import FileUploadParser
import docx
import re
import textract

def content_analysis(str_content,crux_word):
    if not crux_word:
        crux_word = ['奖励', '嘉奖', '补贴', '补助', '补偿', '兜底', '财政支持', '减免', '资助', '税费', '税收', '纳税', '纳税额', '缴纳入库地方级税收',
                     '纳税地方贡献额留成', '纳税地方留成', '地方税收贡献额', '税收地方留成', '地方财政留成', '财政贡献地方留成', '实际形成地方财力', '当地财力贡献', '当地财力实际贡献',
                     '项目库', '名录库', '备选库', '侯选库', '资格库', '储备库', '培育库', '目录库', '供应商库', '中介机构库', '代理机构库', '承包商库', '预选库',
                     '目录', '黑名单', '红名单', '排行榜', '排名榜', '注册', '登记', '备案', '年检', '年报', '监制', '认定', '认可', '检验', '监测', '审定',
                     '指定', '配号', '复检', '复审', '换证', '设立分支机构', '设立法人机构', '设立子公司', '设立分公司', '设有分支机构', '设有子公司', '设有分公司',
                     '常驻机构', '服务网点', '售后机构', '售后服务机构', '迁出', '迁离', '外迁', '搬离', '搬迁', '转出', '追回', '返还', '退还', '收回', '追缴',
                     '承诺', '推荐', '优先给予', '优先支持', '优先扶持', '优先安排', '优先推荐', '优先采购', '优先选用', '优先使用', '优先考虑', '优先选择', '优先选购',
                     '重点支持', '重点扶持', '重点培育', '政策支持', '领军企业', '龙头企业', '知名', '骨干', '国有企业', '民营企业', '限上企业', '规上企业', '本土企业',
                     '本地企业', '本土的企业', '本地的企业', '当地企业', '当地的企业', '外地企业', '外地机构', '区内', '市内', '县内', '本地经营业绩', '本地项目经验',
                     '前置条件', '前置程序', '参考价格', '价格补贴', '统一价格', '本地经营面积', '本地办公面积', '本地固定场所', '本地缴纳社会保险', '减免行政事业性收费',
                     '减免公积金', '减免政府性基金', '缓征行政事业性收费', '缓征公积金', '缓征政府性基金', '停征行政事业性收费', '停征公积金', '停征政府性基金', '减免土地出让金',
                     '减免社会保险费用', '缓征社会保险费用', '缴纳保证金', '以现金形式']

    start_time = datetime.now()
    n = 1
    data_content = []
    serial_list = []
    mapping_list = []
    str_content_text = str_content.split('\n')
    str_content_text = [s_i.strip() for s_i in str_content_text if s_i]

    for i in str_content_text:  # 循环文本的每一行的数据
        i = i.strip()
        if len(i) < 1:
            continue
        n += 1

        def m_i(q, i_data):
            # print(q,i_data)
            if q in mapping_list:
                # print('qwq')

                a_index = mapping_list.index(q)
                mapping_list[:a_index - 1].append(q)
                del serial_list[a_index:]
                serial_list.append(i_data)
            else:
                mapping_list.append(q)
                serial_list.append(i_data)
            # print(serial_list)
            # print()

        i_sign = ""
        if re.findall('^(第.*?章) ', i):  # 匹配    第一章
            # print(i)
            m_i('A', i)
            i_sign = "A"
        elif re.findall('^(第.*?条) ', i):  # 匹配    第一条
            # print(i)
            m_i('B', i)
            i_sign = "B"
        elif re.findall('^([一二三四五六七八九十]+、)', i):  # 匹配     一、
            # print(i)
            m_i('C', i)
            i_sign = "C"
        elif re.findall('^(（[一二三四五六七八九十]+）)', i):  # 匹配    （一）
            # print(i)
            m_i('D', i)
            i_sign = "D"
        elif re.findall('^(\d+\.)', i):  # 匹配     1.
            # print(i)
            m_i('E', i)
            i_sign = "E"
        elif re.findall('^(（\d+）)', i):  # 匹配    （1）
            # print(i)
            m_i('F', i)
            i_sign = "F"
        for c_word in crux_word:
            # if c_word in i:
            sum_i = i.count(c_word)
            if sum_i:

                # print(serial_list,[c_word]+serial_list+[i])
                if i not in serial_list:
                    data_content.append([c_word] + serial_list + [i] + [str(n)] + [str(sum_i)])
                else:
                    data_content.append([c_word] + serial_list + [str(n)] + [str(sum_i)])
                # print(data_content[-1][-3], '--------------', str_content_text.index(i))

                # 段尾获取
                end_text = ''

                for d_c in str_content_text[str_content_text.index(i):]:  # 从当前段落往后循环

                    i_sign1 = ""
                    if re.findall('^(第.*?章) ', d_c):  # 匹配    第一章
                        i_sign1 = "A"
                    elif re.findall('^(第.*?条) ', d_c):  # 匹配    第一条
                        i_sign1 = "B"
                    elif re.findall('^([一二三四五六七八九十]+)', d_c):  # 匹配     一、

                        i_sign1 = "C"
                    elif re.findall('^(（[一二三四五六七八九十]+）)', d_c):  # 匹配    （一）
                        i_sign1 = "D"
                    elif re.findall('^(\d+\.)', d_c):  # 匹配     1.

                        i_sign1 = "E"
                    elif re.findall('^(（\d+）)', d_c):  # 匹配    （1）
                        i_sign1 = "F"
                    else:
                        i_sign1 = "G"
                    # print(mapping_list, i_sign, i_sign1)

                    if i_sign1 == 'G':
                        # end_text = d_c
                        data_content[-1].append(d_c)
                        break
                    if i_sign1 != i_sign:
                        break
    for q in data_content:
        print(q)

    print(datetime.now() - start_time)
    return data_content


def insert_db(data,f_id,f_url,f_name):
    start_time = datetime.now()
    print(start_time)
    with connections['gpjz'].cursor() as cursor:
        for i in data:
            try:
                if i[-1].isalnum():

                    sql = r"""insert into content_analysis_upload(content_id,
                        content_title,
                        content_download_url,
                        word_crux,
                        word_count,
                        word_page,
                        content_appear,
                        content_first,
                        content_end     
                    ) values ("{}","{}","{}","{}",{},"{}","{}","{}","{}")""".format(f_id, f_name, f_url, i[0], i[-1], i[-2], i[-3], i[1:-3], '0')
                    print(sql)
                else:
                    sql = r"""insert into content_analysis_upload(
                        content_id,
                        content_title,
                        content_download_url,
                        word_crux,
                        word_count,
                        word_page,
                        content_appear,
                        content_first,
                        content_end) values ("{}","{}","{}","{}",{},"{}","{}","{}","{}")""".format(f_id, f_name, f_url, i[0], i[-2], i[-3], i[-4], i[1:-4], i[-1])
                    print(sql)

                # 执行sql语句
                cursor.execute(sql)
                # 提交到数据库执行
                connections['gpjz'].commit()
            except Exception as e:
                # 如果发生错误则回滚
                connections['gpjz'].rollback()
                print(e,'这里报错啦！',i)
                break
        print(datetime.now()-start_time)


class Downloads(APIView):
    def get(self, request):
        data = {
            "error_code": 400,
            "params": {
                "error": ""
            }
        }
        id = request.GET.get('id')
        print(id)
        if not id:
            data['params']['error'] = '接口参数缺失！'
            return Response(data=data)
        a = models.File.objects.filter(name=id).first()
        print(a)

        if not a:
            data['error_code'] = 401
            data['params']['error'] = '接口参数错误！'
            return Response(data=data)
        f = a.file
        f_name = str(f).replace('file/','')
        file_path = MEDIA_ROOT.joinpath(str(f))
        file = open(str(file_path), 'rb')
        response =StreamingHttpResponse(file)
        response['Content-Type'] = 'application/octet-stream'
        response['Content-Disposition'] = 'attachment;filename="{}"'.format(escape_uri_path(f_name))  # 这是文件的简单描述，注意写法就是这个固定的写法
        return response


class WordCountView(APIView):
    def get(self, request):

        return Response('no get!!!')

    def post(self, request):
        url_ip = 'http://49.4.31.249/'
        data = {
            "error_code": 400,
            "params": {
                "error": ""
            }
        }
        file_obj = request.FILES.get('file')
        print('file_obj:',file_obj)

        crux_word = request.POST.get('crux_word')
        if crux_word:

            try:
                crux_word = eval(crux_word)
            except:
                crux_word = crux_word.split(',')

        if not file_obj:
            data['params']['error'] = '接口参数缺失！'

            return Response(data=data)

        name_file = models.File.objects.create(
            name=uuid.uuid1(),
            file=file_obj,
        )  # 自动就会将文件上传到我们配置的img文件夹中
        c = name_file.file                  # c: file/南平市人民政府关于印发南平市进一步促进总部经济发展若干措施的通知_M1f3Dz4.docx
        print('c:',c)
        file_path = MEDIA_ROOT.joinpath(str(c))  # 保存文件的路径
        if str(file_obj).split('.')[1] == 'pdf':
            w_content = self.pdf_get(file_path)

        elif str(file_obj).split('.')[1] == 'docx':
            w_content = self.word_get(file_path)
        elif str(file_obj).split('.')[1] == 'doc':
            print(file_path,type(file_path))
            w_content = self.word_get_doc(file_path)
            print(w_content)
        else:
            data['error_code'] = 402
            data['params']['error'] = '文件格式错误！拒绝上传！'
            os.remove(file_path)
            models.File.objects.filter(name=name_file.name).delete()
            return Response(data=data, status=402)
        data = content_analysis(w_content, crux_word)
        # insert_db(data,name_file.name,'{}api/downloads/?id={}'.format(url_ip, name_file.name), str(file_obj))
        print("name_file.name:", name_file.name)
        data_list = []

        for d_i in data:
            print(d_i)
            dict_data = {
                'content_id': str(name_file.name),
                'file_name': str(file_obj),
                'download_url': '{}api/downloads/?id={}'.format(url_ip, name_file.name),
                'word_crux': d_i[0]
            }

            if d_i[-1].isalnum():

                dict_data['word_count'] = d_i[-1]
                dict_data['word_page'] = d_i[-2]
                dict_data['content_appear'] = d_i[-3]
                dict_data['content_first'] = d_i[1:-3]
                dict_data['content_end'] = '0'
            else:
                dict_data['word_count'] = d_i[-2]
                dict_data['word_page'] = d_i[-3]
                dict_data['content_appear'] = d_i[-4]
                dict_data['content_first'] = d_i[1:-4]
                dict_data['content_end'] = d_i[-1]

            data_list.append(dict_data)

        content_data = {
            "error_code": 200,
            "data": data_list

        }

        return Response(status=200, data=content_data, content_type=json)

    def pdf_get(self, pdf_path):
        ee_time = datetime.now()
        d_data = {}
        page_content = ''
        # 内容提取，使用 pdfplumber 打开 PDF，用于提取文本
        with pdfplumber.open(pdf_path) as pdf_file:
            # 使用 PyPDF2 打开 PDF 用于提取图片
            pdf_image_reader = PyPDF2.PdfFileReader(open(pdf_path, "rb"))
            d_data['sum_page'] = pdf_image_reader.getNumPages()  # 获取总页数
            print("开始执行---")
            # len(pdf.pages)为PDF文档页数，一页页解析
            for i in range(len(pdf_file.pages)):

                page_text = pdf_file.pages[i].extract_text()
                if not page_text:
                    continue
                # page.extract_text()函数即读取文本内容
                page_content += page_text

        print("文本内容获取所用时间:", datetime.now() - ee_time)
        return page_content

    def word_get_doc(self, filename):
        f_name = str(filename)
        print(f_name)
        try:
            text = textract.process(f_name)
            text = text.decode()
            return text
        except Exception as err:
            print(err)

    def word_get(self,filename):
        document = docx.Document(filename)
        print(len(document.paragraphs))                     # 段落长度
        content = []
        for i in range(len(document.paragraphs)):
            content.append(document.paragraphs[i].text)

        tables = document.tables  # 获取文件中的表格集
        qwq = []  # word文档里面表格内容
        for t_i in tables:
            for r_j in t_i.rows:
                for z in r_j.cells:
                    if qwq:
                        if z.text not in qwq:
                            qwq.append(z.text)
                    else:
                        qwq.append(z.text)

        word_content = '\n'.join(qwq + content)
        return word_content


class DecimalEncoder(json.JSONEncoder):
    def default(self, o):
        if isinstance(o, decimal.Decimal):
            return float(o)
        elif isinstance(o, dt.date):
            return str(o)
        super(DecimalEncoder, self).default(o)


def index(requests):
    return render(requests, 'index.html')


def choice_database(request, d_name, f_number=1, p_number=100, ):
    print(request.method)
    with connections[d_name].cursor() as cursor:

        # 查看所有表
        cursor.execute("show tables")
        # cursor.execute("SELECT foo FROM bar WHERE baz = %s", [self.baz])
        row = cursor.fetchall()

        t_name = row[0][0]

        if request.method == 'POST':
            req = json.loads(request.body)
            f_number = req.get('f_number')  # 当前页码
            p_number = req.get('p_number')  # 页码数据条数
            t_name = req.get('t_name')  # 数据库表名
            if type(f_number) != int or type(p_number) != int:
                return HttpResponse(status=400, content='参数错误！')
            print(f_number, p_number)
        print(t_name, '*******')
        # 查询有多少数据   默认为第一张表
        cursor.execute("select count(1) as c from %s" % t_name)  # 查询有多少条数据
        # cursor.execute("SELECT foo FROM bar WHERE baz = %s", [self.baz])
        count_t = cursor.fetchone()
        print(count_t[0])
        print(t_name)

        # 查询表字段名及注释 默认为第一张表
        sql_t_name = "select column_name,column_comment from INFORMATION_SCHEMA.Columns where table_name='{}'".format(
            t_name)
        cursor.execute(sql_t_name)  # 查询有多少条数据
        tables_data = cursor.fetchall()

        # 去除 无注释的表字段 ,并生成字典
        tables_name_dict = {}
        for name in tables_data:
            if name[1]:
                tables_name_dict[name[0]] = name[1]
        print(tables_name_dict)

        # 查询表数据
        sql_t_data = "select {} from {} limit {},{}".format(str(list(tables_name_dict)).replace('[', '').
                                                            replace(']', '').replace("'", ''), t_name,
                                                            (f_number - 1) * p_number, p_number)

        cursor.execute(sql_t_data)
        t_data = cursor.fetchall()
        print(sql_t_data)
        tableData = []
        for i in t_data:
            d_2 = {}
            n = 0
            for d in i:
                d_2[str(n)] = d
                n += 1
            tableData.append(d_2)

        print(tableData)
        print('***********')

    response_data = {
        'tables': [i[0] for i in row],
        'count_t': count_t[0],
        'tables_name': list(tables_name_dict.values()),
        "name": 'qwq',
        'tables_data': tableData,
    }
    return response_data


def datatable(request, ):
    print(request.method)

    response_data = choice_database(request, 'analysis1')
    print(response_data)
    if request.method == 'POST':
        return HttpResponse(json.dumps(response_data))

    else:

        return render(request, 'table_data.html', response_data)


def datatable1(request):
    print(request.method)
    sss = datetime.now()
    if request.method == 'POST':

        req = json.loads(request.body)
        d_name = req.get('database_name')  # 当前页码
        t_name = req.get('t_name')
        print(d_name)
        print(t_name)

        with connections[d_name].cursor() as cursor:
            # 查看所有表
            cursor.execute("show tables")
            row = cursor.fetchall()
            cursor.execute("select * from dw_listen_js_table_name")
            tables_name = cursor.fetchall()
            d_tables_name = dict(tables_name)

        response_data = {
            'tables': [{i[0]: d_tables_name[i[0]]} for i in row if i[0] in d_tables_name],
        }

        if t_name:
            response_data = choice_database(request, d_name)
            # for i, t in enumerate(response_data['tables_data']):
            #     print(type(response_data['tables_data'][i]['10']))
            #     print(type(response_data['tables_data'][i]['15']))
            #     response_data['tables_data'][i]['10'] = str(t['10'])
            # #     response_data['tables_data'][i]['11'] = str(t['11'])
            # #     response_data['tables_data'][i]['12'] = str(t['12'])
            #
            #     response_data['tables_data'][i]['15'] = str(t['15'])
        print(response_data)

        print("所用时间:", datetime.now() - sss)
        return HttpResponse(json.dumps(response_data, cls=DecimalEncoder))
    else:
        response_data = {
            'database_name': database_name,
        }
        print("所用时间:", datetime.now() - sss)
        return render(request, 'mid_data.html', response_data)


def chart_table(request, f_number=1, p_number=100, ):
    with connections['analysis1'].cursor() as cursor:

        # 查看所有表
        cursor.execute("show tables")
        # cursor.execute("SELECT foo FROM bar WHERE baz = %s", [self.baz])
        row = cursor.fetchall()

        t_name = row[0][0]

        if request.method == 'POST':

            req = json.loads(request.body)
            t_name = req.get('t_name')  # 数据库表名
            f_number = req.get('f_number')  # 当前页码
            p_number = req.get('p_number')  # 页码数据条数
            if type(f_number) != int or type(p_number) != int:
                return HttpResponse(status=400, content='参数错误！')

        # 查询有多少数据   默认为第一张表
        cursor.execute("select count(1) as c  from %s" % t_name)  # 查询有多少条数据
        # cursor.execute("SELECT foo FROM bar WHERE baz = %s", [self.baz])
        count_t = cursor.fetchone()
        print('count_t:', count_t)
        # 查询表字段名及注释 默认为第一张表
        sql_t_name = "select column_name,column_comment from INFORMATION_SCHEMA.Columns where table_name='{}'".format(
            t_name)
        cursor.execute(sql_t_name)  # 查询有多少条数据
        tables_data = cursor.fetchall()

        # 去除 无注释的表字段 ,并生成字典
        tables_name_dict = {}
        for name in tables_data:
            if name[1]:
                tables_name_dict[name[1]] = name[0]

        # 查询表数据
        sql_t_data = "select {} from {} limit {},{}".format(str(list(tables_name_dict.values())).replace('[', '').
                                                            replace(']', '').replace("'", ''), t_name,
                                                            (f_number - 1) * p_number, p_number)

        cursor.execute(sql_t_data)

        t_data = cursor.fetchall()

        tableData = []
        for i in t_data:
            d_2 = {}
            n = 0
            for d in i:
                d_2[str(n)] = d
                n += 1
            tableData.append(d_2)

    response_data = {
        'tables': [i[0] for i in row],
        'count_t': count_t[0],
        'tables_name': tables_name_dict,
        "name": 'qwq',
        'tables_data': tableData,
    }

    print(response_data)  # 反回的数据
    if request.method == 'POST':
        return HttpResponse(json.dumps(response_data))
    else:

        return render(request, 'table_write.html', response_data)


def show_chart(request):
    """
    接受前端先选择的数据库表字段、可视化图形的选择，
    return: json格式pyecharts模块的需要的数据

    """

    option = {
        '1': '求和',
        '2': '求均值',
        '3': '求最小值',
        '4': '求最大值',
        '5': '数值列求和',
    }

    chart_option = {

        '1': '柱状图',
        '2': '堆积柱状图',
        '3': '百分比堆积柱状图',
        '4': '条形图',
        '5': '堆积条形图',
        '6': '百分比堆积条形图',
        '7': '折线图',
        '8': '堆积折线图',
        '9': '百分比堆积折线图',
        '10': '饼图',
        '11': '环图',
        '12': '复合饼图',
        '13': '复合条饼图',
        '14': '双轴图',
        '15': '面积图',
        '16': '堆积面积图',
        '17': '百分比堆积面积图',
        '18': '散点图',
        '19': '词云',
        '20': '雷达图',
        '21': '指标卡',
        '22': '漏斗图',
        '23': '计量图',
        '24': '箱线图',
        '25': '数据地图',
        '26': 'GIS地图',
        '27': '瀑布图',
        '28': 'GIS地图',
        '29': '旭日图',
        '30': '直方图',
        '31': '帕累托图',
        '32': '桑基图',
        '33': '弦图',
        '34': '矩形树图',
        '35': '甘特图',
    }
    two_charts_choice = {
        1: "right",
        2: "left",
        3: "right",
        4: "left",
    }

    # 获取请求携带的数据
    req = json.loads(request.body)
    print("req:", req)

    t_name = req.get('t_name')  # 表名
    legend_v = req.get('legend_v')  # 图例   以及计量方法
    legend_x = req.get('legend_x')  # 分类轴 x轴
    chart_value = req.get('chart_value')  # 可视化的图形序号
    sort_value = req.get('sort_value')  # 是否排序
    two_charts = req.get('two_charts')  # 是否排序
    ciyun_value = req.get('ciyun_value')  # 词云模板
    font_value = req.get('font_value')  # 词云字体
    ciyun_txt = req.get('ciyun_txt')  # 词云文本
    completion_rate = req.get('completion_rate')  # 计量图的完成率
    print("******", ciyun_value, font_value, ciyun_txt)
    print('two_charts:', two_charts)
    if chart_value == '14':
        legend_v = two_charts

    #  判断是否传参
    if chart_value == '19' and (not ciyun_value or not font_value):
        return HttpResponse(status=400, content='没有词云参数！')
    if not t_name or not legend_v or not legend_x:
        return HttpResponse(status=400, content='参数错误！')

    # 漏斗图的数据必须排序
    if chart_value == '21':
        sort_value = True

    # try:
    with connections['analysis1'].cursor() as cursor:
        # 需要表名      x分类轴字段
        sql_t_name = "select column_name,column_comment from INFORMATION_SCHEMA.Columns where table_name='{}'". \
            format(t_name)
        cursor.execute(sql_t_name)  # 查询表字段和注释
        tables_data = cursor.fetchall()

        # 去除 无注释的表字段 ,并生成字典
        tables_name_dict = {}
        for name in tables_data:
            if name[1]:
                tables_name_dict[name[1]] = name[0]

        # 字段注释 转字段名
        legend_x = tables_name_dict[legend_x]  # 分类轴x轴 字段名称  字段注释 ==> 分类id

        yaxis_dict = {}  # 图例分类列表
        print("legend_v:", legend_v)
        for s in legend_v:

            t_data_get_name = tables_name_dict[s[0]]
            if s[1] == '1' or s[1] == 1:  # 求和
                t_data_get = 'select {},round(sum({}),0) as r from {} GROUP BY {}'.format(legend_x,
                                                                                          t_data_get_name,
                                                                                          t_name, legend_x)

            elif s[1] == '2' or s[1] == 2:  # 求平均值
                t_data_get = 'select {},round(avg({}),0) as r  from {} GROUP BY {}'.format(legend_x,
                                                                                           t_data_get_name,
                                                                                           t_name, legend_x)

            elif s[1] == '3' or s[1] == 3:  # 求最小值
                t_data_get = 'select {},round(min({}),0) as r from {} GROUP BY {}'.format(legend_x, t_data_get_name,
                                                                                          t_name, legend_x)

            elif s[1] == '4' or s[1] == 4:  # 求最大值
                t_data_get = 'select {},round(max({}),0) as r from {} GROUP BY {}'.format(legend_x, t_data_get_name,
                                                                                          t_name, legend_x)
            elif s[1] == '5' or s[1] == 5:  # 数值列求和
                t_data_get = 'select {},round(count({}),0) as r from {} GROUP BY {}'.format(legend_x, t_data_get_name,
                                                                                            t_name, legend_x)
            cursor.execute(t_data_get)
            tables_data = cursor.fetchall()
            if two_charts:
                # print(two_charts, '********************')

                yaxis_dict[str(s[2]) + '.' + s[0] + option[str(s[1])]] = [i[1] for i in tables_data]
            else:
                yaxis_dict[s[0] + option[str(s[1])]] = [i[1] for i in tables_data]

        # x轴聚合之后的列表
        legend_x = 'select {} from {} GROUP BY {}'.format(legend_x, t_name, legend_x)
        cursor.execute(legend_x)
        legend_x_list = cursor.fetchall()  # x轴数组数据

    # except Exception as e:
    #     print(e)
    #     return HttpResponse(status=400, content='参数错误！')
    print('chart_value:', chart_value)

    print("yaxis_dict______>", yaxis_dict)
    y_list_s = list(yaxis_dict.values())
    sum_list = list(np.array(y_list_s).sum(axis=0, dtype=int))  # 二维数组每列求和
    sum_l = []
    for i in y_list_s:
        d = [{"value": k, "percent": round(k / sum_list[i_index], 2)} for i_index, k in enumerate(i)]
        sum_l.append(d)

    legend_x_list = [i[0] for i in legend_x_list]  # x轴 列表数据

    print("legend_x_list:", legend_x_list)  # x轴 列表数据
    print("y_list_s:", y_list_s, len(y_list_s))  # 所有图裂列表数据 (y轴)

    # 排序处理
    if sort_value:
        sort_d = {1: False, 2: True}
        aa = list(zip(legend_x_list, *y_list_s))
        s = sorted(aa, key=lambda function: function[1], reverse=sort_d[sort_value])  # 排序
        s = list(zip(*s))
        legend_x_list = s[0]  # 重新定义x轴
        for i, v in enumerate(yaxis_dict):
            yaxis_dict[v] = s[i + 1]

    # chart1 柱状图
    if chart_value == '1':

        c = Bar(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
        # 可视化图形 生成
        c = c.add_xaxis(legend_x_list)

        for i in yaxis_dict:
            c = c.add_yaxis(i, yaxis_dict[i])
        c = c.set_series_opts(
            label_opts=opts.LabelOpts(position="inside")
        )
        c = c.set_global_opts(title_opts=opts.TitleOpts(title="Bar-基本示例", subtitle="我是副标题"),
                              datazoom_opts=[opts.DataZoomOpts(), opts.DataZoomOpts(type_="inside")],
                              # visualmap_opts=[opts.VisualMapOpts()],
                              toolbox_opts=[opts.ToolboxOpts(
                                  is_show=True,
                                  orient="vertical",
                                  pos_left="90%")],
                              )
        c = c.dump_options_with_quotes()
        # print(c)
        return JsonResponse(json.loads(c))

    # 堆积柱状图
    elif chart_value == '2':
        # 可视化图形 生成
        c = Bar(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
        c = c.add_xaxis(legend_x_list)
        for i in yaxis_dict:
            c = c.add_yaxis(i, yaxis_dict[i], stack="stack1")
        c = c.set_series_opts(
            label_opts=opts.LabelOpts(position="inside")
        )
        c = c.set_global_opts(
            title_opts=opts.TitleOpts(title="堆积柱状图 ", subtitle="我是副标题"),
            datazoom_opts=[opts.DataZoomOpts(), opts.DataZoomOpts(type_="inside")],  # 区域缩放配置项
            # visualmap_opts=[opts.VisualMapOpts()],                                      # 视觉映射配置项
            toolbox_opts=[opts.ToolboxOpts(
                is_show=True,
                orient="vertical",
                pos_left="90%"
            )],  # 工具类
        )

        c = c.dump_options_with_quotes()

        return JsonResponse(json.loads(c))
    # 百分比堆积柱状图
    elif chart_value == '3' or chart_value == '6':

        c = Bar(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
        c = c.add_xaxis(legend_x_list)

        for i, d in enumerate(yaxis_dict):
            y_list = [i['percent'] for i in sum_l[i]]
            c = c.add_yaxis(d, y_axis=y_list, stack="stack1", category_gap="50%")

        c = c.set_series_opts(
            label_opts=opts.LabelOpts(
                position="inside",
                formatter=JsCode(
                    "function(x){return Number(x.data * 100).toFixed() + '%';}"
                    # "function(x){return Number(x.data.percent * 100).toFixed() + '%';}"
                ),
            )
        )
        if chart_value == '6':
            c = c.reversal_axis()
            c = c.set_global_opts(
                title_opts=opts.TitleOpts(title="百分比堆积柱状图 ", subtitle="我是副标题"),
                datazoom_opts=[opts.DataZoomOpts(orient='vertical', type_='slider'), ],
                # visualmap_opts=[opts.VisualMapOpts()],
                toolbox_opts=[opts.ToolboxOpts(
                    is_show=True,
                    orient="vertical",
                    pos_left="90%")],  # 工具类
            )
        else:
            c = c.set_global_opts(
                title_opts=opts.TitleOpts(title="百分比堆积柱状图 ", subtitle="我是副标题"),

                datazoom_opts=[opts.DataZoomOpts(), opts.DataZoomOpts(type_="inside")],
                # visualmap_opts=[opts.VisualMapOpts()],
                toolbox_opts=[opts.ToolboxOpts(
                    is_show=True,
                    orient="vertical",
                    pos_left="90%")],  # 工具类
            )

        c = c.dump_options_with_quotes()
        # print(c)
        return JsonResponse(json.loads(c))

    # 条形图
    elif chart_value == '4':

        c = Bar(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
        # 可视化图形 生成
        c = c.add_xaxis(legend_x_list)

        for i in yaxis_dict:
            c = c.add_yaxis(i, yaxis_dict[i])
        c = c.reversal_axis()
        c = c.set_series_opts(
            label_opts=opts.LabelOpts(position="inside")
        )
        c = c.set_global_opts(

            title_opts=opts.TitleOpts(title="Bar-基本示例", subtitle="我是副标题"),
            datazoom_opts=[opts.DataZoomOpts(orient='vertical', type_='slider'), ],
            # visualmap_opts=[opts.VisualMapOpts()],
            toolbox_opts=[
                opts.ToolboxOpts(
                    is_show=True,
                    orient="vertical",
                    pos_left="90%"
                )
            ],

        )
        c = c.dump_options_with_quotes()

        # print(c)

        return JsonResponse(json.loads(c))

    # 堆积条形图
    elif chart_value == '5':
        # 可视化图形 生成
        c = Bar(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
        c = c.add_xaxis(legend_x_list)

        for i in yaxis_dict:
            c = c.add_yaxis(i, yaxis_dict[i], stack="stack1")
        c = c.reversal_axis()
        c = c.set_series_opts(
            label_opts=opts.LabelOpts(position='inside'),
        )
        c = c.set_global_opts(
            title_opts=opts.TitleOpts(title="堆积条形图 ", subtitle="我是副标题"),
            datazoom_opts=[opts.DataZoomOpts(orient='vertical', type_='slider'), ],  # 区域缩放配置项
            # visualmap_opts=[opts.VisualMapOpts()],                                      # 视觉映射配置项
            toolbox_opts=[opts.ToolboxOpts(
                is_show=True,
                orient="vertical",
                pos_left="90%"
            )],  # 工具类
        )

        c = c.dump_options_with_quotes()

        return JsonResponse(json.loads(c))

    # 折线图
    elif chart_value == '7' or chart_value == '8':
        c = Line(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))

        c = c.add_xaxis(legend_x_list)

        if chart_value == '8':
            for i in yaxis_dict:
                c = c.add_yaxis(i, yaxis_dict[i], is_smooth=True, stack="stack1")
                c = c.set_global_opts(
                    title_opts=opts.TitleOpts(title="堆积折线图"),
                    datazoom_opts=[opts.DataZoomOpts(), opts.DataZoomOpts(type_="inside")],
                    toolbox_opts=[opts.ToolboxOpts(
                        is_show=True,
                        orient="vertical",
                        pos_left="90%")],  # 工具类
                )
        else:
            for i in yaxis_dict:
                c = c.add_yaxis(i, yaxis_dict[i], is_smooth=True)

            c = c.set_global_opts(
                title_opts=opts.TitleOpts(title="普通折线图"),
                datazoom_opts=[opts.DataZoomOpts(), opts.DataZoomOpts(type_="inside")],
                toolbox_opts=[opts.ToolboxOpts(
                    is_show=True,
                    orient="vertical",
                    pos_left="90%")],  # 工具类
            )
        c = c.dump_options_with_quotes()

        return JsonResponse(json.loads(c))

    # 百分比堆积折线图
    elif chart_value == '9':

        c = Line(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
        c = c.add_xaxis(legend_x_list)
        print()
        for i, d in enumerate(yaxis_dict):
            y_list = [i['percent'] for i in sum_l[i]]
            c = c.add_yaxis(d, y_axis=y_list, stack="stack1", is_smooth=True)

        c = c.set_series_opts(
            label_opts=opts.LabelOpts(
                position="inside",
                formatter=JsCode(
                    "function(x){return Number(x.data[1] * 100).toFixed() + '%';}"
                ),
            )
        )
        c = c.set_global_opts(
            title_opts=opts.TitleOpts(title="百分比堆积折线图"),
            datazoom_opts=[opts.DataZoomOpts(), opts.DataZoomOpts(type_="inside")],
            toolbox_opts=[opts.ToolboxOpts(
                is_show=True,
                orient="vertical",
                pos_left="90%")],  # 工具类
        )
        c = c.dump_options_with_quotes()

        return JsonResponse(json.loads(c))

    # 饼图 和 环形图
    elif chart_value == '10' or chart_value == '11':
        pie_list = list(zip(legend_x_list, list(yaxis_dict.values())[0]))
        pie_list = [list(i) for i in pie_list]
        c = Pie(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
        if chart_value == '11':

            c = c.add("", pie_list, radius=["40%", "75%"], )
        else:

            c = c.add("", pie_list, )
        c = c.set_global_opts(
            title_opts=opts.TitleOpts(title="Pie-基本示例"),
            legend_opts=opts.LegendOpts(type_="scroll", pos_top="20%", pos_left="2%", orient="vertical"),
            toolbox_opts=[opts.ToolboxOpts(
                is_show=True,
                orient="vertical",
                pos_top="10%",
                pos_left="90%")],  # 工具类
        )
        c = c.set_series_opts(
            label_opts=opts.LabelOpts(formatter="{b}: {c}  —— {d}%"),

        )
        c = c.dump_options_with_quotes()

        return JsonResponse(json.loads(c))
    # 双轴图
    elif chart_value == '14':
        b = Bar(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
        b = b.add_xaxis(xaxis_data=legend_x_list)
        c = Line(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
        c = c.add_xaxis(legend_x_list)
        two_yd = list(yaxis_dict.keys())

        for i, k in enumerate(two_charts):
            # print( list(yaxis_dict.keys()))
            yd_index = two_yd[i]
            if k[2] < 2:
                b = b.add_yaxis(
                    yd_index, yaxis_dict[yd_index],
                    label_opts=opts.LabelOpts(is_show=False),
                )
            else:
                c = c.add_yaxis(
                    yd_index, yaxis_dict[yd_index],
                    yaxis_index=1,
                    z=10,
                    label_opts=opts.LabelOpts(is_show=False),
                )
        b = b.extend_axis(
            yaxis=opts.AxisOpts(
                # name="温度",
                type_="value",
                position="right",
                # min_=0,
                # max_=25,
                # interval=5,

            )
        )
        b = b.set_series_opts(
            label_opts=opts.LabelOpts(position="inside")
        )
        b = b.set_global_opts(
            title_opts=opts.TitleOpts(title="Bar-基本示例", subtitle="我是副标题"),
            datazoom_opts=[opts.DataZoomOpts(), opts.DataZoomOpts(type_="inside")],

            # yaxis_opts=opts.AxisOpts(
            #     type_="value",
            #     interval=50,
            #     # position="left",
            #     # axistick_opts=opts.AxisTickOpts(is_show=True),
            #     # splitline_opts=opts.SplitLineOpts(is_show=True),
            #     ),
        )

        overlap_1 = b.overlap(c).dump_options_with_quotes()

        return JsonResponse(json.loads(overlap_1))

    # 面积图 与 堆积面积图
    elif chart_value == '15' or chart_value == '16':

        c = Line(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
        c = c.add_xaxis(legend_x_list)
        if chart_value == '16':
            for i in yaxis_dict:
                c = c.add_yaxis(i, yaxis_dict[i], is_smooth=True, stack="stack1")
                c = c.set_series_opts(
                    areastyle_opts=opts.AreaStyleOpts(opacity=0.5),
                    label_opts=opts.LabelOpts(is_show=False),
                )
                c = c.set_global_opts(
                    title_opts=opts.TitleOpts(title="Line-堆积面积图"),
                    datazoom_opts=[opts.DataZoomOpts(), opts.DataZoomOpts(type_="inside")],
                    xaxis_opts=opts.AxisOpts(
                        axistick_opts=opts.AxisTickOpts(is_align_with_label=True),
                        is_scale=False,
                        boundary_gap=False,
                    ),
                    toolbox_opts=[opts.ToolboxOpts(
                        is_show=True,
                        orient="vertical",
                        pos_left="90%")],  # 工具类
                )
        else:
            for i in yaxis_dict:
                c = c.add_yaxis(i, yaxis_dict[i], is_smooth=True)
            c = c.set_series_opts(
                areastyle_opts=opts.AreaStyleOpts(opacity=0.5),
                label_opts=opts.LabelOpts(is_show=False),
            )
            c = c.set_global_opts(
                title_opts=opts.TitleOpts(title="Line-面积图"),
                xaxis_opts=opts.AxisOpts(
                    axistick_opts=opts.AxisTickOpts(is_align_with_label=True),
                    is_scale=False,
                    boundary_gap=False,
                ),
                datazoom_opts=[opts.DataZoomOpts(), opts.DataZoomOpts(type_="inside")],
                toolbox_opts=[opts.ToolboxOpts(
                    is_show=True,
                    orient="vertical",
                    pos_left="90%")],  # 工具类
            )
        c = c.dump_options_with_quotes()

        return JsonResponse(json.loads(c))

    # 百分比堆积面积图
    elif chart_value == '17':

        c = Line(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
        c = c.add_xaxis(legend_x_list)

        for i, d in enumerate(yaxis_dict):
            y_list = [i['percent'] for i in sum_l[i]]
            c = c.add_yaxis(d, y_axis=y_list, stack="stack1", is_smooth=True)
        c = c.set_series_opts(
            areastyle_opts=opts.AreaStyleOpts(opacity=0.5),
            label_opts=opts.LabelOpts(
                is_show=False,
                position="inside",
                formatter=JsCode(
                    "function(x){return Number(x.data[1] * 100).toFixed() + '%';}"
                ),
            )
        )
        c = c.set_global_opts(
            title_opts=opts.TitleOpts(title="百分比堆积面积图"),
            xaxis_opts=opts.AxisOpts(
                axistick_opts=opts.AxisTickOpts(is_align_with_label=True),
                is_scale=False,
                boundary_gap=False,
            ),
            datazoom_opts=[opts.DataZoomOpts(), opts.DataZoomOpts(type_="inside")],
            toolbox_opts=[opts.ToolboxOpts(
                is_show=True,
                orient="vertical",
                pos_left="90%")],  # 工具类
        )
        c = c.dump_options_with_quotes()

        return JsonResponse(json.loads(c))
        # return HttpResponse(a.render_embed())
        # return HttpResponse(json.loads(line_base()))
        # c.render_embed()
        # a.render("gauge_base.html")

        # return HttpResponse('ok!')

    # 散点图
    elif chart_value == '18':
        c = Scatter(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
        c = c.add_xaxis(legend_x_list)
        max_y = 0
        for i in yaxis_dict:
            i_max = max(yaxis_dict[i])
            if i_max > max_y:
                max_y = i_max
            c = c.add_yaxis(i, yaxis_dict[i])

        c = c.set_global_opts(
            title_opts=opts.TitleOpts(title="Scatter-多维度数据"),
            visualmap_opts=opts.VisualMapOpts(
                type_="color", max_=max_y + 5, min_=0, dimension=1
            ),
            toolbox_opts=[opts.ToolboxOpts(
                is_show=True,
                orient="vertical",
                pos_left="90%")],
        )
        c = c.dump_options_with_quotes()

        return JsonResponse(json.loads(c))

    # 词云图
    elif chart_value == '19':
        if not ciyun_txt:
            ciyun_txt = '无'
        # 词云分词处理
        print(ciyun_txt)
        pair_text = jieba.posseg.lcut(ciyun_txt)
        pair_no = ['x', 'p', 'uj']
        c_text_list = [tuple(i)[0] for i in pair_text if tuple(i)[1] not in pair_no]
        s = Counter(c_text_list).items()

        s_sort = sorted(s, key=lambda qwq: qwq[1], reverse=True)
        print(s_sort)
        # print(s_sort)
        # 获取词云模板文件目录路径
        # img_list = os.listdir(STATIC_URL_TEMPLATE)
        # print(STATIC_URL_TEMPLATE)
        # print(font_value,ciyun_value)
        mask_iamge = STATIC_URL_TEMPLATE.joinpath(ciyun_value)  # 获取词云模板图片路径
        # print("mask_iamge;",STATIC_URL_TEMPLATE.joinpath(ciyun_value))

        c = WordCloud(init_opts=opts.InitOpts(bg_color='white'))
        c = c.add("", s_sort, word_size_range=[20, 100],
                  is_draw_out_of_bound=False,
                  textstyle_opts=opts.TextStyleOpts(font_family=font_value),
                  mask_image=mask_iamge
                  )
        c = c.set_global_opts(title_opts=opts.TitleOpts(title="WordCloud"))
        c = c.dump_options_with_quotes()
        # print(c)
        return JsonResponse(json.loads(c))

        # except Exception as e:
        #     print(e)

    # 雷达图
    elif chart_value == '20':

        c_schema = []  # 轴名称 外围标签
        for i, k in enumerate(legend_x_list[:6]):
            c_schema.append({'name': k})

        c = Radar(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
        c = c.add_schema(schema=c_schema, shape="polygon", )

        # 数据线的颜色库
        color_list = ['#fa5a5a', '#f0d264', '#82c8a0', '#7fccde', '#6698cb', '#cb99c5', '#ff0000', '#eb4310',
                      '#f6941d', '#fbb417', '#ffff00', '#cdd541', '#99cc33', '#3f9337', '#219167', '#239676', '#24998d',
                      '#1f9baa', '#0080ff', '#3366cc', '#333399', '#003366', '#800080', '#a1488e', '#c71585', '#bd2158']
        for i in yaxis_dict:
            c = c.add(i, [yaxis_dict[i][:6]], color=random.choice(color_list))

        c = c.set_series_opts(label_opts=opts.LabelOpts(is_show=False))
        c = c.set_global_opts(
            title_opts=opts.TitleOpts(title="雷达图"),
            toolbox_opts=[opts.ToolboxOpts(
                is_show=True,
                orient="vertical",
                pos_left="90%")],
        )
        c = c.dump_options_with_quotes()

        return JsonResponse(json.loads(c))

        # except Exception as e:
        #     print(e)

    # 指标卡
    elif chart_value == "21":
        pass

    # 漏斗图
    elif chart_value == '22':
        pie_list = list(zip(legend_x_list, list(yaxis_dict.values())[0]))
        pie_list = [list(i) for i in pie_list]

        c = Funnel(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))

        x_data = ["展现量", "点击量", "访问量", "咨询量", "订单量"]
        y_data = [100, 79, 50, 31, 21]
        x_data1 = ["点击率", "访问率", "咨询率", "最终转化率", ]
        y_data1 = [98, 62, 39, 21]
        data1 = [[x_data1[i], y_data1[i]] for i in range(len(x_data1))]
        data = [[x_data[i], y_data[i]] for i in range(len(x_data))]
        # c = c.add(
        #         series_name="",
        #         data_pair=data,
        #         gap=2,
        #         tooltip_opts=opts.TooltipOpts(trigger="item", formatter="{a} <br/>{b} : {c}%"),
        #         label_opts=opts.LabelOpts(is_show=True, position="inside"),
        #         itemstyle_opts=opts.ItemStyleOpts(border_color="#fff", border_width=1),
        #     )
        c = c.add("商品", data,
                  tooltip_opts=opts.TooltipOpts(trigger="item", formatter="{a} <br/>{b} : {c}%"),
                  # label_opts=opts.LabelOpts(is_show=True, position="inside"),
                  itemstyle_opts=opts.ItemStyleOpts(border_color="#fff", border_width=1),
                  )
        c = c.set_global_opts(
            title_opts=opts.TitleOpts(title="漏斗图", subtitle=" "),
            toolbox_opts=[opts.ToolboxOpts(
                is_show=True,
                orient="vertical",
                pos_left="95%")],
        )
        c = c.dump_options_with_quotes()

        return JsonResponse(json.loads(c))

    # 计量图
    elif chart_value == "23":
        try:
            c = Gauge(init_opts=opts.InitOpts(theme=ThemeType.LIGHT, bg_color='white'))
            c = c.add(
                series_name="业务指标",
                data_pair=[["完成率", eval(completion_rate)]],
                detail_label_opts=opts.GaugeDetailOpts(
                    offset_center=[0, '40%'],
                    font_size=30,
                    formatter="{value}%"

                )
            )
            c = c.set_global_opts(
                legend_opts=opts.LegendOpts(is_show=False),
                tooltip_opts=opts.TooltipOpts(is_show=True, formatter="{a} <br/>{b} : {c}%"),
                toolbox_opts=[opts.ToolboxOpts(
                    is_show=True,
                    orient="vertical",
                    pos_left="90%")],
            )
            c = c.dump_options_with_quotes()
            # print(c)

            return JsonResponse(json.loads(c))
        except Exception as err:
            print(err)
            return HttpResponse(status=400, content="填写的完成率格式不对！")

    # 箱线图
    elif chart_value == "24":
        return HttpResponse(status=400, content="开发中！！！")

    # 地图面积图(颜色深浅)
    elif chart_value == "25":

        c = Map()
        a = [[k, v] for k, v in zip([t[:-3] if '自治区' in t else t[:-1] for t in legend_x_list],
                                    list(yaxis_dict.values())[0])]
        max_y = 0
        for i in yaxis_dict:
            i_max = max(yaxis_dict[i])
            if i_max > max_y:
                max_y = i_max
        c = c.add("数据地图", a, "china")
        c = c.set_global_opts(
            title_opts=opts.TitleOpts(title="Map-地图"),
            visualmap_opts=opts.VisualMapOpts(
                min_=0, max_=max_y,
            )
        )
        c = c.dump_options_with_quotes()
        # print(c)
        return JsonResponse(json.loads(c))

    #  # 地图面积图(热力图)
    elif chart_value == "26":

        #
        a = [[k, v] for k, v in zip([t[:-3] if '自治区' in t else t[0:-1] for t in
                                     ['新疆自治区' if e == '新疆维吾尔自治区' else e for e in legend_x_list]],
                                    list(yaxis_dict.values())[0])]
        c = Geo()
        c = c.add_schema(maptype="china")
        c = c.add(
            "geo",
            a,
            type_=ChartType.HEATMAP,

        )
        c = c.set_series_opts(

            label_opts=opts.LabelOpts(is_show=True))
        # 定义最大颜色值
        max_y = 0
        for i in yaxis_dict:
            i_max = max(yaxis_dict[i])
            if i_max > max_y:
                max_y = i_max
        c = c.set_global_opts(
            visualmap_opts=opts.VisualMapOpts(
                min_=0, max_=max_y,
            ),
            title_opts=opts.TitleOpts(title="Geo-HeatMap"),
        )

        c = c.dump_options_with_quotes()
        # print(c)
        return JsonResponse(json.loads(c))

        return HttpResponse(status=400, content="开发中！！！")

    # 计量图
    elif chart_value == "27":
        return HttpResponse(status=400, content="开发中！！！")

    # 计量图
    elif chart_value == "23":
        return HttpResponse(status=400, content="开发中！！！")

    return HttpResponse(status=400, content='字段错误无法生成图形！')


def word_count(request):
    a = ['奖励', '补贴', '补助', '资助', '税费', '税收', '纳税', '纳税额', '缴纳入库地方级税收', '纳税地方贡献额留成', '纳税地方留成', '地方税收贡献额', '税收地方留成',
         '地方财政留成',
         '财政贡献地方留成', '项目库', '名录库', '备选库', '资格库', '储备库', '培育库', '目录库', '目录', '注册', '登记', '设立分支机构', '设立法人机构', '设立子公司',
         '设有分支机构', '设有法人机构', '设有子公司', '迁出', '迁离', '外迁', '转出', '追回', '返还', '退还', '收回', '追缴', '承诺', '优先给予', '优先支持',
         '优先扶持',
         '优先安排', '优先推荐', '优先采购', '优先选用', '优先使用', '优先考虑', '优先选择', '优先选购', '重点支持', '重点扶持', '领军企业', '龙头', '知名', '骨干', '推荐',
         '本土', '本地', '区内']
    a1 = ['奖励1', '补贴1', '补助1', '资助1', '税费1', '税收1', '纳税1', '纳税额1', '缴纳入库地方级税收1', '纳税地方贡献额留成1', '纳税地方留成1', '地方税收贡献额1',
          '税收地方留成1', '地方财政留成1', '财政贡献地方留成1', '项目库1', '名录库1', '备选库1', '资格库1', '储备库1', '培育库1', '目录库1', '目录1', '注册1',
          '登记1', '设立分支机构1', '设立法人机构1', '设立子公司1', '设有分支机构1', '设有法人机构1', '设有子公司1', '迁出1', '迁离1', '外迁1', '转出1', '追回1',
          '返还1', '退还1', '收回1', '追缴1', '承诺1', '优先给予1', '优先支持1', '优先扶持1', '优先安排1', '优先推荐1', '优先采购1', '优先选用1', '优先使用1',
          '优先考虑1', '优先选择1', '优先选购1', '重点支持1', '重点扶持1', '领军企业1', '龙头1', '知名1', '骨干1', '推荐1', '本土1', '本地1', '区内1']
    word_data_name = {
        'a': a,
        'a1': a1,
    }
    """
    接受上传文件，反回词数
    """
    if request.method == "GET":
        # 高频词汇库
        response_data = [{'label': "a", 'value': 'a'}, {'label': "a1", 'value': 'a1'}]

        return render(request, 'word_count.html', {'data': response_data})
    else:

        file_obj = request.FILES.get('file')
        # 获取请求携带的数据
        print('file_obj:', file_obj)

        if not file_obj:
            req = json.loads(request.body)

            word_data = req.get('word_v')  # 表名
            word_list = []
            for i, v in enumerate(word_data_name[word_data]):
                word_list.append({'id': i, 'name': v})
            return HttpResponse(json.dumps(word_list))

        # file_obj = request.FILES.get('file')
        print(request.POST)
        word_str = request.POST.get('word_str')  # 用户处理后的高频词汇
        # print(word_str)
        print('word_str:', word_str, type(word_str))
        if word_str and word_str != '[]':
            print('word_str:', word_str)

            a = eval(word_str)

        name_file = models.File.objects.create(
            name=uuid.uuid1(),
            file=file_obj,
        )  # 自动就会将文件上传到我们配置的img文件夹中
        c = name_file.file

        file_path = MEDIA_ROOT.joinpath(str(c))  # 保存文件的路径
        if str(file_obj).split('.')[1] == 'pdf':

            d_data = extract_content(file_path, a, str(file_obj))

        elif str(file_obj).split('.')[1] == 'docx':
            d_data = extract_content_docx(file_path, a, str(file_obj))

        else:

            os.remove(file_path)
            return HttpResponse(status=400, content="文件格式错误！拒绝上传")
        d_list = []
        id_n = 1
        for i in d_data:
            if i == 'sum_page':
                continue

            d_list.append({"id": id_n, 'file_name': str(file_obj), "name": i, **d_data[i]})
            id_n += 1
        # print(d_list)
        id_n = 1
        for d_list_i in d_list:
            nn = id_n * 1000
            for children_i in d_list_i['children']:
                children_i['id'] = nn
                children_i['name'] = d_list_i['name']
                children_i['file_name'] = str(file_obj)
                nn += 1
            id_n += 1
        print(d_list)
        return HttpResponse(json.dumps(d_list))


def extract_content_docx(path, a, f_name):
    s_time = datetime.now()
    print(s_time)
    d_a = {}

    document = Document(path)  # 读入文件

    tables = document.tables  # 获取文件中的表格集

    qwq = []  # word文档里面表格内容
    for t_i in tables:
        for r_j in t_i.rows:
            for z in r_j.cells:
                if qwq:
                    if z.text not in qwq:
                        qwq.append(z.text)
                else:
                    qwq.append(z.text)

    row_content = '\n'.join(qwq)

    for i in range(len(document.paragraphs)):
        content = document.paragraphs[i].text

        multiple(i, a, content, d_a, f_name)

    for t in a:
        t_count = row_content.count(t)
        if t_count:
            t_content = '      \n'.join([q_i for q_i in qwq if t in q_i])

            if d_a.get(t, False):
                if d_a[t].get('content'):

                    d_a[t]['content'] += t_content
                else:
                    d_a[t]['content'] = t_content
                d_a[t]['str_sum'] += t_count
                d_a[t]['str_page'] += f'、表格'
                d_a[t]['children'] += [{'str_page': '表格', 'str_sum': t_count, 'content': t_content, }]
            else:
                d_a[t] = {'str_page': '表格', 'str_sum': t_count, 'content': t_content,
                          'children': [{'str_page': '表格', 'str_sum': t_count, 'content': t_content, }]}
    # print(row_content)

    print(datetime.now() - s_time)

    return d_a


def extract_content(pdf_path, a, file_obj):
    ee_time = datetime.now()
    d_data = {}
    # 内容提取，使用 pdfplumber 打开 PDF，用于提取文本
    with pdfplumber.open(pdf_path) as pdf_file:
        # 使用 PyPDF2 打开 PDF 用于提取图片
        pdf_image_reader = PyPDF2.PdfFileReader(open(pdf_path, "rb"))
        d_data['sum_page'] = pdf_image_reader.getNumPages()  # 获取总页数

        # len(pdf.pages)为PDF文档页数，一页页解析
        for i in range(len(pdf_file.pages)):
            page_text = pdf_file.pages[i]
            # page.extract_text()函数即读取文本内容
            page_content = page_text.extract_text()
            multiple(i, a, page_content, d_data, file_obj)

    print("ee:", datetime.now() - ee_time)
    # print(d_data)

    return d_data


def multiple(i, a, page_content, d_data, f_name):
    """
    i:页数
    a:关键字库
    page_content:段落内容
    d_data:反回数据的字典
    """

    for j in a:

        page_content = str(page_content)
        str_sum = page_content.count(j)  # 每一页字符出现的次数
        if str_sum:

            if d_data.get(j, False):
                # d_data[j]['file_name'] = f_name
                if d_data[j].get('content'):
                    d_data[j]['content'] += page_content
                else:
                    d_data[j]['content'] = page_content

                d_data[j]['str_sum'] += str_sum

                d_data[j]['str_page'] += f'、{i}'

                d_data[j]['children'] += [{'str_page': i, 'str_sum': str_sum, 'content': page_content, }]

            else:
                d_data[j] = {
                    # 'file_name': f_name,
                    'content': page_content,
                    'str_page': str(i),
                    'str_sum': str_sum,
                    'children': [{'str_page': i, 'str_sum': str_sum, 'content': page_content}]
                }

    return d_data
