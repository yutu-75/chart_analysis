import docx
from docx import Document  # 导入库

path = "仙游县人民政府关于印发仙仙游县促进建筑业发展壮大的实施意见的通知.docx"  # 文件路径
document = Document(path)       # 读入文件

tables = document.tables #获取文件中的表格集
a = []
for i in tables:
    for j in i.rows:
        for z in j.cells:
            # print(z.text)
            if a:

                if z.text not in a:

                        a.append(z.text)
            else:
                a.append(z.text)



print(a)








doc = Document(path) #filename为word文档

#获取文档中的表格
doc.tables  #获取文档的表格个数 len(doc.tables)
# print(len(doc.tables))
#读取第1个表格
tb1=doc.tables[0]

#获取第一个表格的行
tb1.rows  #获取表格的行数len(tb1.rows)
# print(len(tb1.rows))
#读取表格的第一行的单元格
row_cells=tb1.rows[3].cells
# print(row_cells[3].text)
#读取第一行所有单元格的内容
n = 1
# for cell in row_cells:
#
#
#     print(cell.text)





